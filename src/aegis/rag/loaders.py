"""
Document Loaders

Load documents from various sources:
- PDF files (clinical reports, policies)
- Word documents (procedures, guidelines)
- Text files (notes, logs)
- HL7 messages
- FHIR resources
"""

from typing import Any, Dict, List, Optional, BinaryIO
from datetime import datetime
from abc import ABC, abstractmethod
from pathlib import Path
import json
import re

import structlog
from pydantic import BaseModel, Field

logger = structlog.get_logger(__name__)


# =============================================================================
# Document Model
# =============================================================================

class LoadedDocument(BaseModel):
    """A loaded document ready for processing."""
    id: str
    content: str
    metadata: Dict[str, Any] = Field(default_factory=dict)
    
    # Source info
    source: str  # File path or URL
    source_type: str  # pdf, docx, txt, hl7, fhir
    
    # Document info
    title: Optional[str] = None
    author: Optional[str] = None
    created_at: Optional[datetime] = None
    
    # Processing info
    loaded_at: datetime = Field(default_factory=datetime.utcnow)
    page_count: int = 1
    char_count: int = 0
    
    def __post_init__(self):
        self.char_count = len(self.content)


# =============================================================================
# Base Loader
# =============================================================================

class DocumentLoader(ABC):
    """Abstract base class for document loaders."""
    
    @abstractmethod
    def load(self, source: str) -> List[LoadedDocument]:
        """Load documents from source."""
        pass
    
    @abstractmethod
    def load_bytes(self, data: bytes, filename: str) -> List[LoadedDocument]:
        """Load documents from bytes."""
        pass
    
    def _generate_id(self, source: str) -> str:
        """Generate document ID from source."""
        import hashlib
        return hashlib.md5(source.encode()).hexdigest()[:16]


# =============================================================================
# PDF Loader
# =============================================================================

class PDFLoader(DocumentLoader):
    """
    Load PDF documents.
    
    Uses PyMuPDF (fitz) for extraction.
    Handles:
    - Text extraction
    - Table detection
    - Image OCR (optional)
    - Metadata extraction
    """
    
    def __init__(self, ocr_enabled: bool = False):
        self.ocr_enabled = ocr_enabled
    
    def load(self, source: str) -> List[LoadedDocument]:
        """Load PDF from file path."""
        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"PDF not found: {source}")
        
        with open(path, "rb") as f:
            return self.load_bytes(f.read(), path.name)
    
    def load_bytes(self, data: bytes, filename: str) -> List[LoadedDocument]:
        """Load PDF from bytes."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.warning("PyMuPDF not installed, using fallback")
            return self._fallback_load(data, filename)
        
        documents = []
        
        try:
            pdf = fitz.open(stream=data, filetype="pdf")
            
            # Extract metadata
            metadata = pdf.metadata or {}
            
            # Extract text from all pages
            full_text = []
            for page_num in range(len(pdf)):
                page = pdf[page_num]
                text = page.get_text()
                full_text.append(text)
            
            content = "\n\n".join(full_text)
            
            doc = LoadedDocument(
                id=self._generate_id(filename),
                content=content,
                source=filename,
                source_type="pdf",
                title=metadata.get("title") or filename,
                author=metadata.get("author"),
                page_count=len(pdf),
                metadata={
                    "creator": metadata.get("creator"),
                    "producer": metadata.get("producer"),
                    "subject": metadata.get("subject"),
                    "keywords": metadata.get("keywords"),
                },
            )
            
            documents.append(doc)
            pdf.close()
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return self._fallback_load(data, filename)
        
        return documents
    
    def _fallback_load(self, data: bytes, filename: str) -> List[LoadedDocument]:
        """Fallback when PyMuPDF is not available."""
        return [LoadedDocument(
            id=self._generate_id(filename),
            content=f"[PDF content from {filename} - install PyMuPDF for extraction]",
            source=filename,
            source_type="pdf",
            title=filename,
        )]


# =============================================================================
# Word Document Loader
# =============================================================================

class DocxLoader(DocumentLoader):
    """
    Load Word documents (.docx).
    
    Uses python-docx for extraction.
    """
    
    def load(self, source: str) -> List[LoadedDocument]:
        """Load DOCX from file path."""
        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {source}")
        
        with open(path, "rb") as f:
            return self.load_bytes(f.read(), path.name)
    
    def load_bytes(self, data: bytes, filename: str) -> List[LoadedDocument]:
        """Load DOCX from bytes."""
        try:
            from docx import Document as DocxDocument
            import io
        except ImportError:
            logger.warning("python-docx not installed")
            return [LoadedDocument(
                id=self._generate_id(filename),
                content=f"[DOCX content - install python-docx]",
                source=filename,
                source_type="docx",
            )]
        
        try:
            doc = DocxDocument(io.BytesIO(data))
            
            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    table_rows.append(row_text)
                tables_text.append("\n".join(table_rows))
            
            content = "\n\n".join(paragraphs)
            if tables_text:
                content += "\n\n[Tables]\n" + "\n\n".join(tables_text)
            
            # Extract metadata
            core_props = doc.core_properties
            
            return [LoadedDocument(
                id=self._generate_id(filename),
                content=content,
                source=filename,
                source_type="docx",
                title=core_props.title or filename,
                author=core_props.author,
                created_at=core_props.created,
                metadata={
                    "subject": core_props.subject,
                    "keywords": core_props.keywords,
                    "category": core_props.category,
                },
            )]
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return [LoadedDocument(
                id=self._generate_id(filename),
                content=f"[DOCX extraction failed: {e}]",
                source=filename,
                source_type="docx",
            )]


# =============================================================================
# Text Loader
# =============================================================================

class TextLoader(DocumentLoader):
    """Load plain text files."""
    
    def __init__(self, encoding: str = "utf-8"):
        self.encoding = encoding
    
    def load(self, source: str) -> List[LoadedDocument]:
        """Load text from file path."""
        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {source}")
        
        with open(path, "r", encoding=self.encoding) as f:
            content = f.read()
        
        return [LoadedDocument(
            id=self._generate_id(source),
            content=content,
            source=source,
            source_type="txt",
            title=path.name,
        )]
    
    def load_bytes(self, data: bytes, filename: str) -> List[LoadedDocument]:
        """Load text from bytes."""
        content = data.decode(self.encoding)
        
        return [LoadedDocument(
            id=self._generate_id(filename),
            content=content,
            source=filename,
            source_type="txt",
            title=filename,
        )]


# =============================================================================
# HL7 Loader
# =============================================================================

class HL7Loader(DocumentLoader):
    """
    Load HL7v2 messages as documents.
    
    Extracts structured content from HL7 messages.
    """
    
    def load(self, source: str) -> List[LoadedDocument]:
        """Load HL7 from file."""
        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {source}")
        
        with open(path, "r") as f:
            content = f.read()
        
        return self._parse_hl7(content, source)
    
    def load_bytes(self, data: bytes, filename: str) -> List[LoadedDocument]:
        """Load HL7 from bytes."""
        content = data.decode("utf-8")
        return self._parse_hl7(content, filename)
    
    def _parse_hl7(self, content: str, source: str) -> List[LoadedDocument]:
        """Parse HL7 message into document."""
        from aegis.integrations.hl7v2 import HL7v2Parser
        
        try:
            message = HL7v2Parser.parse(content)
            
            # Convert to readable text
            text_parts = [
                f"HL7 Message: {message.message_type}^{message.trigger_event}",
                f"Control ID: {message.control_id}",
                f"Patient ID: {message.patient_id}",
            ]
            
            family, given = message.patient_name
            if family or given:
                text_parts.append(f"Patient Name: {given} {family}")
            
            # Add segment details
            for segment in message.segments:
                text_parts.append(f"\n{segment.segment_id}: {segment.raw[:200]}")
            
            return [LoadedDocument(
                id=self._generate_id(source),
                content="\n".join(text_parts),
                source=source,
                source_type="hl7",
                title=f"HL7 {message.message_type}^{message.trigger_event}",
                metadata={
                    "message_type": message.message_type,
                    "trigger_event": message.trigger_event,
                    "patient_id": message.patient_id,
                },
            )]
            
        except Exception as e:
            logger.error(f"HL7 parsing failed: {e}")
            return [LoadedDocument(
                id=self._generate_id(source),
                content=content,
                source=source,
                source_type="hl7",
            )]


# =============================================================================
# FHIR Loader
# =============================================================================

class FHIRLoader(DocumentLoader):
    """
    Load FHIR resources as documents.
    
    Converts FHIR JSON to readable text.
    """
    
    def load(self, source: str) -> List[LoadedDocument]:
        """Load FHIR from file."""
        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {source}")
        
        with open(path, "r") as f:
            content = f.read()
        
        return self._parse_fhir(content, source)
    
    def load_bytes(self, data: bytes, filename: str) -> List[LoadedDocument]:
        """Load FHIR from bytes."""
        content = data.decode("utf-8")
        return self._parse_fhir(content, filename)
    
    def _parse_fhir(self, content: str, source: str) -> List[LoadedDocument]:
        """Parse FHIR resource into document."""
        try:
            data = json.loads(content)
            resource_type = data.get("resourceType", "Unknown")
            
            # Handle bundles
            if resource_type == "Bundle":
                documents = []
                for entry in data.get("entry", []):
                    resource = entry.get("resource", {})
                    doc = self._resource_to_document(resource, source)
                    if doc:
                        documents.append(doc)
                return documents if documents else [self._resource_to_document(data, source)]
            else:
                return [self._resource_to_document(data, source)]
                
        except json.JSONDecodeError as e:
            logger.error(f"FHIR JSON parsing failed: {e}")
            return [LoadedDocument(
                id=self._generate_id(source),
                content=content,
                source=source,
                source_type="fhir",
            )]
    
    def _resource_to_document(self, resource: dict, source: str) -> LoadedDocument:
        """Convert FHIR resource to document."""
        resource_type = resource.get("resourceType", "Unknown")
        resource_id = resource.get("id", "unknown")
        
        # Build readable text based on resource type
        text_parts = [f"FHIR {resource_type} (ID: {resource_id})"]
        
        if resource_type == "Patient":
            name = resource.get("name", [{}])[0]
            text_parts.append(f"Name: {name.get('given', [''])[0]} {name.get('family', '')}")
            text_parts.append(f"Birth Date: {resource.get('birthDate', 'N/A')}")
            text_parts.append(f"Gender: {resource.get('gender', 'N/A')}")
            
        elif resource_type == "Condition":
            code = resource.get("code", {}).get("coding", [{}])[0]
            text_parts.append(f"Condition: {code.get('display', code.get('code', 'N/A'))}")
            text_parts.append(f"Status: {resource.get('clinicalStatus', {}).get('coding', [{}])[0].get('code', 'N/A')}")
            
        elif resource_type == "Observation":
            code = resource.get("code", {}).get("coding", [{}])[0]
            text_parts.append(f"Observation: {code.get('display', code.get('code', 'N/A'))}")
            value = resource.get("valueQuantity", {})
            if value:
                text_parts.append(f"Value: {value.get('value')} {value.get('unit', '')}")
        
        # Add raw JSON for complete info
        text_parts.append(f"\nRaw: {json.dumps(resource, indent=2)[:1000]}")
        
        return LoadedDocument(
            id=self._generate_id(f"{source}_{resource_id}"),
            content="\n".join(text_parts),
            source=source,
            source_type="fhir",
            title=f"{resource_type}/{resource_id}",
            metadata={
                "resource_type": resource_type,
                "resource_id": resource_id,
            },
        )


# =============================================================================
# Loader Factory
# =============================================================================

class DocumentLoaderFactory:
    """Factory for creating document loaders."""
    
    LOADERS = {
        ".pdf": PDFLoader,
        ".docx": DocxLoader,
        ".doc": DocxLoader,
        ".txt": TextLoader,
        ".md": TextLoader,
        ".hl7": HL7Loader,
        ".json": FHIRLoader,
    }
    
    @classmethod
    def get_loader(cls, file_path: str) -> DocumentLoader:
        """Get appropriate loader for file type."""
        ext = Path(file_path).suffix.lower()
        loader_class = cls.LOADERS.get(ext, TextLoader)
        return loader_class()
    
    @classmethod
    def load(cls, file_path: str) -> List[LoadedDocument]:
        """Load a document using appropriate loader."""
        loader = cls.get_loader(file_path)
        return loader.load(file_path)
